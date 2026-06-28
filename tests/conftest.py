import zipfile
from pathlib import Path
import pytest
import fitz  # pymupdf
from openpyxl import Workbook
from docx import Document


@pytest.fixture
def born_digital_pdf(tmp_path) -> Path:
    p = tmp_path / "born.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Renewable energy zone cost benefit analysis")
    doc.save(p)
    doc.close()
    return p


@pytest.fixture
def image_only_pdf(tmp_path) -> Path:
    p = tmp_path / "drawing.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.draw_rect(fitz.Rect(100, 100, 300, 300), fill=(0, 0, 0))
    doc.save(p)
    doc.close()
    return p


@pytest.fixture
def mixed_pdf(tmp_path) -> Path:
    p = tmp_path / "mixed.pdf"
    doc = fitz.open()
    t = doc.new_page()
    t.insert_text((72, 72), "Page one has real text content here")
    g = doc.new_page()
    g.draw_rect(fitz.Rect(100, 100, 300, 300), fill=(0, 0, 0))
    doc.save(p)
    doc.close()
    return p


@pytest.fixture
def simple_docx(tmp_path) -> Path:
    p = tmp_path / "report.docx"
    d = Document()
    d.add_heading("Business Case", level=0)
    d.add_paragraph("This project delivers value.")
    d.save(p)
    return p


@pytest.fixture
def simple_xlsx(tmp_path) -> Path:
    p = tmp_path / "data.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Inputs"
    ws["A1"] = "Region"
    ws["B1"] = "Capacity"
    ws["A2"] = "NSW"
    ws["B2"] = 1200
    wb.save(p)
    return p


@pytest.fixture
def messy_xlsx(tmp_path) -> Path:
    p = tmp_path / "messy.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Title spanning header"
    ws.merge_cells("A1:C1")
    ws["A2"] = "Region"
    ws["B2"] = "Capacity"
    wb.save(p)
    return p


@pytest.fixture
def simple_kmz(tmp_path) -> Path:
    p = tmp_path / "zones.kmz"
    kml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<kml xmlns="http://www.opengis.net/kml/2.2"><Document>'
        '<name>Indicative REZ Boundaries 2025</name>'
        '<Placemark><name>Zone A</name></Placemark>'
        '<Placemark><name>Zone B</name></Placemark>'
        '</Document></kml>'
    )
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("doc.kml", kml)
    return p
